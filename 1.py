import os
import json
import sys
from bs4 import BeautifulSoup
from docx import Document
import chardet

scp ='''
<script src="acim.js"></script>
'''
bodyH='''
<!DOCTYPE html>
<head>
<meta charset=UTF-8>
<link rel="stylesheet" href="style.css">
<style>

</style>
<script src="jquery-3.7.1.min.js"></script>


<title>
</title>
</head>
<body>

'''
bodyF='''
</body>
</html>
'''
print(bodyH)
		


input_num=sys.argv[1]
js = json.load(open('json/'+input_num+'.json'))
#print(js)
if js['volumeId'] == 'text':
	if js['titleInfo'].get('title') :
		print('<div id="title">%s</div>'%js['titleInfo']['title'])
	else:
		print('<div id="title">%s</div>'%js['seoTitle'])
	print(js['bodyHtml'])
	c_num = js['humanId'].split('-')[1]
	if len(js['humanId'].split('-'))<3:
		s_num = 0
	else:
		s_num = js['humanId'].split('-')[3]
	rpath = ''
	for f in os.listdir('miracle_r'):
		if f.split('.')[0] == c_num:
			for f2 in os.listdir('miracle_r/'+f):
				if f2.split('.')[0] == s_num:
					rpath='miracle_r'+'/'+f+'/'+f2
	#print(rpath)
	#text_num = js['annotation'][2:]

	'''
	r
	'''
	det_f = open(rpath,'rb')
	cha = chardet.detect(det_f.read())
	det_f.close()
	unc = 'GB18030'
	#print(cha)
	if cha['encoding'] == 'GB2312':
		unc = 'GB18030'
	else:
		unc=cha['encoding']
	#print(unc)
	
	with open(rpath,'r',encoding=unc) as f:

		content = f.read()
		content = content.replace('？','。')
		content = content.replace('！','。')
		content = content.replace('：','。')
		lst_c_num = []
		for i in range(1,30):
			t = str(i)+'、'
			d = content.find(t)
			if d > 0:
				lst_c_num.append(d)
		#print(lst_c_num)
		lst_c = []
		lst_c_2=[]
		for i in range(len(lst_c_num)):
			if i < len(lst_c_num)-1:
				lst_c.append(content[lst_c_num[i]:lst_c_num[i+1]])
			else:
				lst_c.append(content[lst_c_num[i]:])
		for i in range(len(lst_c)):
			lst_c_2.append([i+1,lst_c[i].strip().split('。')])

		for i in lst_c_2:
			#print(i)
			for j in range(len(i[1])):
				if len(i[1][j])>1:
					print('<div class="popup" tid="%s">%s</div>'%('d'+input_num+'#'+str(i[0])+':'+str(j+1),i[1][j].strip()+'。'))
					#print('u'+input_num+'#'+str(i[0])+':'+str(j+1)+'=='+i[1][j].strip()+'。')

	'''
	xzc
	'''
	for i in os.listdir('xzc/text/'+c_num):
		#print(i[0:len(c_num+'-'+s_num+' ')])
		if i[0:len(c_num+'-'+s_num+' ')] == c_num+'-'+s_num+' ':
			#print(i)
			doc = Document('xzc/text/'+c_num+'/'+i)
			content = ''
			text = []
			for p in doc.paragraphs:
				text.append(p.text)
			content = '\n'.join(text)
			content = content.replace('？','。')
			content = content.replace('！','。')
			content = content.replace('：','。')

			#print(content)
			lst_c_num = []
			for i in range(1,30):
				t = str(i)+'.'
				d = content.find(t)
				if d >= 0:
					lst_c_num.append(d)
			#print(lst_c_num)
			lst_c = []
			lst_c_2=[]
			for i in range(len(lst_c_num)):
				if i < len(lst_c_num)-1:
					lst_c.append(content[lst_c_num[i]:lst_c_num[i+1]])
				else:
					lst_c.append(content[lst_c_num[i]:])
			for i in range(len(lst_c)):
				lst_c_2.append([i+1,lst_c[i].strip().split('。')])

			for i in lst_c_2:
				#print(i)
				for j in range(len(i[1])):
					if len(i[1][j])>1:
						print('<div class="popup" tid="%s">%s</div>'%('u'+input_num+'#'+str(i[0])+':'+str(j+1),i[1][j].strip()+'。'))
						#print('d'+input_num+'#'+str(i[0])+':'+str(j+1)+'=='+i[1][j].strip()+'。')
if js['volumeId'] == 'workbook':
	if js['titleInfo'].get('title') :
		print('<div id="title">%s</div>'%js['titleInfo']['title'])
	else:
		print('<div id="title">%s</div>'%js['seoTitle'])
	print(js['bodyHtml'])
	lesson_num = js['annotation'][2:]

	'''
	r
	'''
	with open('miracle_l/'+lesson_num+'.txt','r',encoding='gb2312') as f:
		content = f.read()
		content = content.replace('？','。')
		content = content.replace('！','。')
		content = content.replace('：','。')
		#content = content.replace('；','。')
		lst_c_num = []
		for i in range(1,30):
			t = str(i)+'、'
			d = content.find(t)
			if d > 0:
				lst_c_num.append(d)
		#print(lst_c_num)
		lst_c = []
		lst_c_2=[]
		for i in range(len(lst_c_num)):
			if i < len(lst_c_num)-1:
				lst_c.append(content[lst_c_num[i]:lst_c_num[i+1]])
			else:
				lst_c.append(content[lst_c_num[i]:])
		for i in range(len(lst_c)):
			lst_c_2.append([i+1,lst_c[i].strip().split('。')])

		for i in lst_c_2:
			#print(i)
			for j in range(len(i[1])):
				if len(i[1][j])>1:
					print('<div class="popup" tid="%s">%s</div>'%('d'+input_num+'#'+str(i[0])+':'+str(j+1),i[1][j].strip()+'。'))
					#print('u'+input_num+'#'+str(i[0])+':'+str(j+1)+'=='+i[1][j].strip()+'。')

	'''
	xzc
	'''
	for i in os.listdir('xzc/workbook/'):
		if i[0:len(lesson_num)+len('課')] == lesson_num+'課':
			doc = Document('xzc/workbook/'+i)
			content = ''
			text = []
			for p in doc.paragraphs:
				text.append(p.text)
			content = '\n'.join(text)
			content = content.replace('？','。')
			content = content.replace('！','。')
			content = content.replace('：','。')
			#content = content.replace('；','。')
			#print(content)
			lst_c_num = []
			for i in range(1,30):
				t = str(i)+'.'
				d = content.find(t)
				if d >= 0:
					lst_c_num.append(d)
			#print(lst_c_num)
			lst_c = []
			lst_c_2=[]
			for i in range(len(lst_c_num)):
				if i < len(lst_c_num)-1:
					lst_c.append(content[lst_c_num[i]:lst_c_num[i+1]])
				else:
					lst_c.append(content[lst_c_num[i]:])
			for i in range(len(lst_c)):
				lst_c_2.append([i+1,lst_c[i].strip().split('。')])

			for i in lst_c_2:
				#print(i)
				for j in range(len(i[1])):
					if len(i[1][j])>1:
						print('<div class="popup" tid="%s">%s</div>'%('u'+input_num+'#'+str(i[0])+':'+str(j+1),i[1][j].strip()+'。'))
						#print('d'+input_num+'#'+str(i[0])+':'+str(j+1)+'=='+i[1][j].strip()+'。')
print('<div display="none" class="popup" id="utext"></div>')
print('<div display="none" class="popup" id="dtext"></div>')
print(scp)
print(bodyF)